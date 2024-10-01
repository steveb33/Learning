import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
from statsmodels.tsa.stattools import adfuller
from docx import Document
from docx.shared import Inches, Pt

# Load data with the updated paths
data = pd.read_csv('/Users/stevenbarnes/Desktop/Dissertation/ALL_DATA.csv')
summary = pd.read_csv('/Users/stevenbarnes/Desktop/Dissertation/Tableau Dashboards/Summary.csv')

# Create a new Word document
doc = Document()
doc.add_heading('Data Description', 0)

# Create the directory if it does not exist
output_dir = "/Users/stevenbarnes/Desktop/Dissertation/Data Desc"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)


# Function to calculate coefficient of variation
def calc_cv(data):
    return np.std(data) / np.mean(data)


# Function to calculate ADF test p-value after cleaning data
def calc_adf_pvalue(series):
    clean_series = series.dropna()  # Remove NaN values
    clean_series = clean_series[np.isfinite(clean_series)]  # Remove Inf values
    if len(clean_series) > 0:  # Ensure there are enough data points for the test
        adf_result = adfuller(clean_series, autolag='AIC')
        return adf_result[1]  # p-value is the second element in the result tuple
    else:
        return np.nan  # Return NaN if the cleaned data is insufficient for the test


# Initialize a counter to keep track of the number of variables per row
counter = 0

# Iterate over each category and sub-category
categories = summary['Category'].unique()
for category in categories:
    sub_categories = summary[summary['Category'] == category]['Sub-Category'].unique()
    for sub_category in sub_categories:
        # Add a new heading for each sub-category
        doc.add_heading(f'{category} - {sub_category}', level=1)

        # Get the variables in the current sub-category
        variables = summary[(summary['Category'] == category) &
                            (summary['Sub-Category'] == sub_category)]['Variable'].values

        for i, var in enumerate(variables):
            if var in data.columns:
                # Plot the data with a figure size of (2.5, 1.5)
                plt.figure(figsize=(2, 1.5))
                plt.plot(data[var])

                # Adjust the title and label font sizes
                plt.title(var, fontsize=6)

                # Adjust the tick labels font size
                plt.xticks(fontsize=6)
                plt.yticks(fontsize=6)

                plt.grid(True)
                plot_path = os.path.join(output_dir, f'{var}.png')
                plt.savefig(plot_path, bbox_inches='tight')  # Ensure tight layout
                plt.close()

                # Add plot to the Word document
                if counter == 0:
                    table_row = doc.add_table(rows=1, cols=4)  # Adjusted for 4 variables per row
                    table_row.autofit = True
                cell = table_row.cell(0, counter)
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(plot_path, width=Inches(1.5))  # Adjusted plot width in Word

                # Calculate statistics
                min_val = data[var].min()
                max_val = data[var].max()
                mean_val = data[var].mean()
                median_val = data[var].median()
                cv_val = calc_cv(data[var])
                adf_pvalue = calc_adf_pvalue(data[var])  # Calculate ADF p-value

                # Add summary statistics table below the figure, transposed and smaller font
                table = cell.add_table(rows=7, cols=2)  # Adjusted for 7 rows now
                table.style = 'Table Grid'

                # Add row labels and values with smaller font size
                table.cell(0, 0).text = 'Statistic'
                table.cell(0, 1).text = 'Value'
                table.cell(1, 0).text = 'Min'
                table.cell(1, 1).text = f'{min_val:.2f}'
                table.cell(2, 0).text = 'Max'
                table.cell(2, 1).text = f'{max_val:.2f}'
                table.cell(3, 0).text = 'Mean'
                table.cell(3, 1).text = f'{mean_val:.2f}'
                table.cell(4, 0).text = 'Median'
                table.cell(4, 1).text = f'{median_val:.2f}'
                table.cell(5, 0).text = 'CV'
                table.cell(5, 1).text = f'{cv_val:.4f}'
                table.cell(6, 0).text = 'ADF p-value'
                table.cell(6, 1).text = f'{adf_pvalue:.4f}'  # Add ADF p-value

                for row in table.rows:
                    for cell in row.cells:
                        cell_font = cell.paragraphs[0].runs[0].font
                        cell_font.size = Pt(6)  # Set the font size to 6 points

                # Increment the counter and check if we need to start a new row
                counter += 1
                if counter >= 4:  # Adjusted to 4 variables per row
                    counter = 0
                    doc.add_paragraph('')  # Add space between rows

        # Reset the counter when a sub-category changes to ensure new rows start correctly
        counter = 0

# Save the Word document
doc_path = os.path.join(output_dir, 'Data_Description.docx')
doc.save(doc_path)

print(f'Document saved at {doc_path}')
